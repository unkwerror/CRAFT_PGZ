"""Извлечение текста из документов ТЗ (PDF/DOCX) для передачи в LLM.

Текст размечается маркерами страниц («===== стр. N =====»), чтобы модель могла
цитировать с номером страницы. Объём ограничен MAX_CHARS — для экономии токенов на
очень больших файлах (лишнее отсекается с пометкой truncated).

Сканы без текстового слоя дают почти пустой текст — помечаем low_text, чтобы
предупредить пользователя (OCR — отдельная задача, здесь не делаем).
"""

from __future__ import annotations

import io
from dataclasses import dataclass

import structlog
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

log = structlog.get_logger()

# ~375k токенов при ~4 симв/токен — влезает в окно Sonnet (1M) с запасом, но
# ограничивает стоимость входа на «очень больших» ТЗ (лишнее отсекаем).
MAX_CHARS = 1_500_000
_LOW_TEXT_CHARS_PER_PAGE = 50  # меньше — вероятно скан без текстового слоя


@dataclass
class ExtractResult:
    text: str  # текст с маркерами страниц
    kind: str  # "pdf" | "docx"
    pages: int  # число страниц (для docx = 0)
    char_count: int
    truncated: bool  # текст обрезан по MAX_CHARS
    low_text: bool  # подозрение на скан (мало текста на страницу)


class UnsupportedDocumentError(Exception):
    """Тип файла не поддерживается для разбора текста (нужен PDF или DOCX)."""


def _is_pdf(filename: str, content_type: str | None) -> bool:
    return filename.lower().endswith(".pdf") or (content_type or "") == "application/pdf"


def _is_docx(filename: str, content_type: str | None) -> bool:
    return filename.lower().endswith(".docx") or (content_type or "").endswith(
        "wordprocessingml.document"
    )


def _is_xlsx(filename: str, content_type: str | None) -> bool:
    return filename.lower().endswith(".xlsx") or (content_type or "").endswith(
        "spreadsheetml.sheet"
    )


def _cap(text: str) -> tuple[str, bool]:
    if len(text) <= MAX_CHARS:
        return text, False
    return text[:MAX_CHARS] + "\n\n[...текст обрезан по лимиту разбора...]", True


def _extract_pdf(data: bytes) -> ExtractResult:
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001 — битая страница не должна валить весь файл
            log.warning("pdf_page_extract_failed", page=i, error=str(exc))
            page_text = ""
        parts.append(f"\n\n===== стр. {i} =====\n{page_text.strip()}")
    pages = len(reader.pages)
    body = "".join(parts).strip()
    text, truncated = _cap(body)
    # Мало текста на страницу -> скорее всего скан (нет текстового слоя). Работает и для
    # 1-страничных: раньше был порог pages>3 и малостраничные сканы не распознавались.
    low_text = len(body) < max(1, pages) * _LOW_TEXT_CHARS_PER_PAGE
    return ExtractResult(text, "pdf", pages, len(text), truncated, low_text)


def _extract_docx(data: bytes) -> ExtractResult:
    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    body = "\n".join(parts).strip()
    text, truncated = _cap(body)
    return ExtractResult(text, "docx", 0, len(text), truncated, low_text=not body)


def _extract_xlsx(data: bytes) -> ExtractResult:
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        parts: list[str] = []
        for ws in wb.worksheets:
            parts.append(f"\n\n===== лист: {ws.title} =====")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
    finally:
        wb.close()
    body = "\n".join(parts).strip()
    text, truncated = _cap(body)
    return ExtractResult(text, "xlsx", 0, len(text), truncated, low_text=not body)


def extract_text(filename: str, content_type: str | None, data: bytes) -> ExtractResult:
    """Извлечь текст из PDF/DOCX/XLSX. Прочие типы -> UnsupportedDocumentError."""
    if _is_pdf(filename, content_type):
        return _extract_pdf(data)
    if _is_docx(filename, content_type):
        return _extract_docx(data)
    if _is_xlsx(filename, content_type):
        return _extract_xlsx(data)
    raise UnsupportedDocumentError(filename)
